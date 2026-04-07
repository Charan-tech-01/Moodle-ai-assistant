"""
response_formatter.py — Document generation for Moodle AI Assistant
=====================================================================
Produces well-formatted output in four formats:
  PDF   → reportlab (rich styling) or fpdf2 fallback
  Excel → openpyxl  (branded tables, alternating rows, auto-width)
  Word  → python-docx (heading styles, tables, bold text)
  Text  → plain passthrough

Install dependencies (run once in your project folder):
  pip install reportlab python-docx openpyxl
"""

from __future__ import annotations

import re
import tempfile
from datetime import datetime
from pathlib import Path

# ── openpyxl (Excel) ──────────────────────────────────────────────────────────
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

# ── python-docx (Word) ────────────────────────────────────────────────────────
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ── PDF library — try reportlab first, fall back gracefully ──────────────────
try:
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm as rl_cm
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    from reportlab.platypus import (
        HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    )
    _HAS_REPORTLAB = True
except ImportError:
    _HAS_REPORTLAB = False

try:
    from fpdf import FPDF
    _HAS_FPDF = True
except ImportError:
    _HAS_FPDF = False


# ── Brand palette ─────────────────────────────────────────────────────────────
B_BLUE    = "#1a2b8c"
B_ACCENT  = "#4361ee"
B_LIGHT   = "#eef1fb"
B_BORDER  = "#d0d5f0"
B_ALT     = "#f5f7ff"
B_DARK    = "#111827"
B_MID     = "#374151"
B_MUTED   = "#6b7280"


# ═══════════════════════════════════════════════════════════════════════════════
#  SHARED HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def _ts() -> str:
    return datetime.now().strftime("%d %b %Y, %I:%M %p")


def _clean_md(text: str) -> str:
    """Strip markdown markers."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*",     r"\1", text)
    text = re.sub(r"__(.+?)__",     r"\1", text)
    text = re.sub(r"`(.+?)`",       r"\1", text)
    return text


def _parse(answer: str) -> list[dict]:
    """Parse LLM answer into typed line objects."""
    out = []
    for raw in answer.splitlines():
        line = raw.rstrip()
        if not line:
            out.append({"type": "blank", "text": ""})
        elif line.startswith("### "):
            out.append({"type": "h3", "text": line[4:].strip()})
        elif line.startswith("## "):
            out.append({"type": "h2", "text": line[3:].strip()})
        elif line.startswith("# "):
            out.append({"type": "h1", "text": line[2:].strip()})
        elif re.match(r"^[-=*_]{3,}$", line.strip()):
            out.append({"type": "hr", "text": ""})
        elif re.match(r"^[-*•]\s+", line):
            out.append({"type": "bullet", "text": re.sub(r"^[-*•]\s+", "", line)})
        elif re.match(r"^\d+[.)]\s+", line):
            out.append({"type": "numbered", "text": re.sub(r"^\d+[.)]\s+", "", line)})
        else:
            out.append({"type": "text", "text": line})
    return out


def _extract_table(answer: str):
    """Return (headers, rows) if a markdown table exists, else (None, None)."""
    tlines = [l.strip() for l in answer.splitlines()
              if re.match(r"^\|.+\|$", l.strip())]
    if len(tlines) < 2:
        return None, None
    data = [l for l in tlines if not re.match(r"^\|[-: |]+\|$", l)]
    if not data:
        return None, None

    def split(l):
        return [c.strip() for c in l.strip("|").split("|")]

    return split(data[0]), [split(l) for l in data[1:]]


# ═══════════════════════════════════════════════════════════════════════════════
#  TEXT
# ═══════════════════════════════════════════════════════════════════════════════

def format_text_response(answer: str) -> str:
    """Strip markdown formatting for clean chat display."""
    text = answer
    text = re.sub(r"[*][*](.+?)[*][*]", r"\1", text)
    text = re.sub(r"[*](.+?)[*]",       r"\1", text)
    text = re.sub(r"__(.+?)__",         r"\1", text)
    text = re.sub(r"`(.+?)`",           r"\1", text)
    text = re.sub(r"^[#]{1,3}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^[-*][*]?[*]?\s+", "- ", text, flags=re.MULTILINE)
    text = re.sub(r"^[-=*_]{3,}$", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def create_text_file(answer: str) -> Path:
    with tempfile.NamedTemporaryFile(
        mode="w", encoding="utf-8",
        suffix=".txt", prefix="moodle_ai_", delete=False,
    ) as f:
        f.write(f"NMIT Moodle AI Assistant\nGenerated: {_ts()}\n{'='*60}\n\n{answer}")
    return Path(f.name)


# ═══════════════════════════════════════════════════════════════════════════════
#  PDF
# ═══════════════════════════════════════════════════════════════════════════════

def create_pdf(answer: str) -> Path:
    if _HAS_REPORTLAB:
        return _pdf_reportlab(answer)
    if _HAS_FPDF:
        return _pdf_fpdf(answer)
    # Last resort: save as plain text with .pdf extension
    return create_text_file(answer).rename(
        create_text_file(answer).with_suffix(".pdf")
    )


# ── reportlab implementation ──────────────────────────────────────────────────

def _pdf_reportlab(answer: str) -> Path:
    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", prefix="moodle_ai_", delete=False)
    tmp.close()

    doc = SimpleDocTemplate(
        tmp.name, pagesize=A4,
        leftMargin=2*rl_cm, rightMargin=2*rl_cm,
        topMargin=2.5*rl_cm, bottomMargin=2.5*rl_cm,
    )

    base = getSampleStyleSheet()

    def S(name, **kw):
        return ParagraphStyle(name, parent=base["Normal"], **kw)

    styles = {
        "title":    S("t",  fontName="Helvetica-Bold", fontSize=18,
                       textColor=rl_colors.HexColor(B_BLUE),
                       spaceAfter=3, leading=22),
        "meta":     S("m",  fontName="Helvetica",      fontSize=9,
                       textColor=rl_colors.HexColor(B_MUTED), spaceAfter=10),
        "h1":       S("h1", fontName="Helvetica-Bold", fontSize=14,
                       textColor=rl_colors.HexColor(B_BLUE),
                       spaceBefore=14, spaceAfter=5, leading=18),
        "h2":       S("h2", fontName="Helvetica-Bold", fontSize=11,
                       textColor=rl_colors.HexColor(B_ACCENT),
                       spaceBefore=10, spaceAfter=4),
        "h3":       S("h3", fontName="Helvetica-BoldOblique", fontSize=10,
                       textColor=rl_colors.HexColor(B_MID),
                       spaceBefore=7, spaceAfter=3),
        "body":     S("b",  fontName="Helvetica", fontSize=10,
                       textColor=rl_colors.HexColor(B_DARK),
                       leading=15, spaceAfter=4),
        "bullet":   S("bl", fontName="Helvetica", fontSize=10,
                       textColor=rl_colors.HexColor(B_DARK),
                       leading=14, spaceAfter=3, leftIndent=16),
        "footer":   S("f",  fontName="Helvetica", fontSize=8,
                       textColor=rl_colors.HexColor(B_MUTED),
                       alignment=TA_RIGHT),
    }

    story = []

    # Header
    story.append(Paragraph("Moodle AI Assistant — NMIT", styles["title"]))
    story.append(Paragraph(f"Generated: {_ts()}", styles["meta"]))
    story.append(HRFlowable(width="100%", thickness=1.5,
                             color=rl_colors.HexColor(B_ACCENT), spaceAfter=12))

    # Table detection
    headers, rows = _extract_table(answer)
    if headers and rows:
        non_tbl = re.sub(r"\|.*\|", "", answer, flags=re.MULTILINE).strip()
        if non_tbl:
            for p in _parse(non_tbl):
                _rl_line(story, p, styles)

        cw = doc.width / max(len(headers), 1)
        tdata = [[h.upper() for h in headers]] + rows
        tbl = Table(tdata, colWidths=[cw]*len(headers), repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND",    (0,0), (-1,0),  rl_colors.HexColor(B_BLUE)),
            ("TEXTCOLOR",     (0,0), (-1,0),  rl_colors.white),
            ("FONTNAME",      (0,0), (-1,0),  "Helvetica-Bold"),
            ("FONTSIZE",      (0,0), (-1,0),  9),
            ("FONTNAME",      (0,1), (-1,-1), "Helvetica"),
            ("FONTSIZE",      (0,1), (-1,-1), 9),
            ("ROWBACKGROUNDS",(0,1), (-1,-1), [rl_colors.white, rl_colors.HexColor(B_ALT)]),
            ("GRID",          (0,0), (-1,-1), 0.5, rl_colors.HexColor(B_BORDER)),
            ("ALIGN",         (0,0), (-1,-1), "LEFT"),
            ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
            ("TOPPADDING",    (0,0), (-1,-1), 5),
            ("BOTTOMPADDING", (0,0), (-1,-1), 5),
            ("LEFTPADDING",   (0,0), (-1,-1), 8),
        ]))
        story.append(tbl)
    else:
        for p in _parse(answer):
            _rl_line(story, p, styles)

    # Footer
    story.append(Spacer(1, 1*rl_cm))
    story.append(HRFlowable(width="100%", thickness=0.5,
                             color=rl_colors.HexColor(B_BORDER)))
    story.append(Paragraph(
        "NMIT Smart Campus · Moodle AI Assistant · Confidential",
        styles["footer"]
    ))

    doc.build(story)
    return Path(tmp.name)


def _rl_line(story, p, styles):
    t, text = p["type"], _clean_md(p["text"])
    if t == "blank":
        story.append(Spacer(1, 5))
    elif t == "hr":
        story.append(HRFlowable(width="100%", thickness=0.5,
                                 color=rl_colors.HexColor(B_BORDER),
                                 spaceBefore=6, spaceAfter=6))
    elif t == "h1":
        story.append(Paragraph(text, styles["h1"]))
    elif t == "h2":
        story.append(Paragraph(text, styles["h2"]))
    elif t == "h3":
        story.append(Paragraph(text, styles["h3"]))
    elif t in ("bullet", "numbered"):
        story.append(Paragraph(f"{'•' if t=='bullet' else '→'} {text}", styles["bullet"]))
    elif text:
        story.append(Paragraph(text, styles["body"]))


# ── fpdf fallback ─────────────────────────────────────────────────────────────

def _pdf_fpdf(answer: str) -> Path:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "NMIT Moodle AI Assistant", ln=True)
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(0, 6, f"Generated: {_ts()}", ln=True)
    pdf.ln(4)
    pdf.set_draw_color(67, 97, 238)
    pdf.set_line_width(0.8)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(6)

    for p in _parse(answer):
        t, text = p["type"], _clean_md(p["text"])
        if t == "blank":
            pdf.ln(3)
        elif t == "hr":
            pdf.set_draw_color(200, 200, 200)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(4)
        elif t in ("h1", "h2"):
            pdf.set_font("Helvetica", "B", 13 if t == "h1" else 11)
            pdf.multi_cell(0, 8, text)
            pdf.set_font("Helvetica", "", 10)
        elif t in ("bullet", "numbered"):
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 6, f"  - {text}")
        elif text:
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 6, text)

    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", prefix="moodle_ai_", delete=False)
    tmp.close()
    pdf.output(tmp.name)
    return Path(tmp.name)


# ═══════════════════════════════════════════════════════════════════════════════
#  EXCEL
# ═══════════════════════════════════════════════════════════════════════════════

def create_excel(answer: str) -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Moodle AI Report"

    HDR_BG  = "1A2B8C"
    HDR_FG  = "FFFFFF"
    ACC_BG  = "4361EE"
    ALT_BG  = "F0F3FF"
    BORDER  = "C7CEEA"
    META_FG = "6B7280"
    TITLE_FG= "1A2B8C"

    def thin_border():
        s = Side(border_style="thin", color=BORDER)
        return Border(left=s, right=s, top=s, bottom=s)

    # ── Title block ───────────────────────────────────────────────────────────
    ws.merge_cells("A1:H1")
    c = ws["A1"]
    c.value     = "NMIT Moodle AI Assistant — Report"
    c.font      = Font(name="Calibri", bold=True, size=16, color=TITLE_FG)
    c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
    ws.row_dimensions[1].height = 34

    ws.merge_cells("A2:H2")
    c = ws["A2"]
    c.value     = f"Generated: {_ts()}"
    c.font      = Font(name="Calibri", size=10, color=META_FG, italic=True)
    c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
    ws.row_dimensions[2].height = 18

    # Accent bar
    for col in range(1, 9):
        ws.cell(row=3, column=col).fill = PatternFill("solid", fgColor=ACC_BG)
    ws.row_dimensions[3].height = 4

    row = 5

    # ── Table detection ───────────────────────────────────────────────────────
    headers, rows = _extract_table(answer)
    if headers and rows:
        non_tbl = re.sub(r"\|.*\|", "", answer, flags=re.MULTILINE).strip()
        if non_tbl:
            row = _xl_text(ws, non_tbl, row, HDR_BG, HDR_FG, ALT_BG)
            row += 1

        ncols = len(headers)
        # Header row
        for ci, h in enumerate(headers, 1):
            c = ws.cell(row=row, column=ci, value=h.upper())
            c.font      = Font(name="Calibri", bold=True, size=11, color=HDR_FG)
            c.fill      = PatternFill("solid", fgColor=HDR_BG)
            c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            c.border    = Border(
                top=Side(border_style="medium", color=HDR_BG),
                bottom=Side(border_style="medium", color=HDR_BG),
                left=Side(border_style="thin", color="FFFFFF"),
                right=Side(border_style="thin", color="FFFFFF"),
            )
        ws.row_dimensions[row].height = 24
        row += 1

        # Data rows
        for ri, data_row in enumerate(rows):
            bg = ALT_BG if ri % 2 else "FFFFFF"
            for ci, val in enumerate(data_row[:ncols], 1):
                c = ws.cell(row=row, column=ci, value=val)
                c.font      = Font(name="Calibri", size=10)
                c.fill      = PatternFill("solid", fgColor=bg)
                c.alignment = Alignment(vertical="center", wrap_text=True, indent=1)
                c.border    = Border(
                    bottom=Side(border_style="thin", color=BORDER),
                    left=Side(border_style="thin",   color=BORDER),
                    right=Side(border_style="thin",  color=BORDER),
                )
            ws.row_dimensions[row].height = 20
            row += 1

        # Auto column widths
        for ci in range(1, ncols + 1):
            col_values = [
                len(str(ws.cell(row=r, column=ci).value or ""))
                for r in range(row - len(rows) - 1, row)
            ]
            ws.column_dimensions[get_column_letter(ci)].width = min(max(col_values, default=10) + 4, 40)
    else:
        row = _xl_text(ws, answer, row, HDR_BG, HDR_FG, ALT_BG)

    # ── Footer ────────────────────────────────────────────────────────────────
    row += 2
    ws.merge_cells(f"A{row}:H{row}")
    c = ws.cell(row=row, column=1, value="NMIT Smart Campus · Moodle AI Assistant · Confidential")
    c.font      = Font(name="Calibri", size=9, color=META_FG, italic=True)
    c.alignment = Alignment(horizontal="right")

    ws.freeze_panes = "A5"

    tmp = tempfile.NamedTemporaryFile(suffix=".xlsx", prefix="moodle_ai_", delete=False)
    tmp.close()
    wb.save(tmp.name)
    return Path(tmp.name)


def _xl_text(ws, answer, start_row, hdr_bg, hdr_fg, alt_bg):
    ACC = "4361EE"
    row = start_row
    n   = 0  # numbered list counter

    for p in _parse(answer):
        t, text = p["type"], _clean_md(p["text"])

        if t == "blank":
            row += 1
            continue

        if t == "hr":
            for col in range(1, 9):
                ws.cell(row=row, column=col).fill = PatternFill("solid", fgColor="D0D5F0")
            ws.row_dimensions[row].height = 3
            row += 1
            continue

        ws.merge_cells(f"A{row}:H{row}")
        c = ws.cell(row=row, column=1)

        if t == "h1":
            c.value     = text.upper()
            c.font      = Font(name="Calibri", bold=True, size=13, color=hdr_fg)
            c.fill      = PatternFill("solid", fgColor=hdr_bg)
            c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
            ws.row_dimensions[row].height = 26
        elif t in ("h2", "h3"):
            c.value     = text
            c.font      = Font(name="Calibri", bold=True, size=11, color=ACC)
            c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
            ws.row_dimensions[row].height = 22
        elif t == "bullet":
            c.value     = f"    •   {text}"
            c.font      = Font(name="Calibri", size=10)
            c.fill      = PatternFill("solid", fgColor=alt_bg)
            c.alignment = Alignment(horizontal="left", vertical="center", indent=2)
            ws.row_dimensions[row].height = 18
        elif t == "numbered":
            n += 1
            c.value     = f"    {n}.   {text}"
            c.font      = Font(name="Calibri", size=10)
            c.fill      = PatternFill("solid", fgColor=alt_bg if n % 2 else "FFFFFF")
            c.alignment = Alignment(horizontal="left", vertical="center", indent=2)
            ws.row_dimensions[row].height = 18
        else:
            c.value     = text
            c.font      = Font(name="Calibri", size=10)
            c.alignment = Alignment(horizontal="left", vertical="center",
                                    wrap_text=True, indent=1)
            ws.row_dimensions[row].height = 18

        c.border = Border(bottom=Side(border_style="thin", color="E5E7EB"))
        row += 1

    ws.column_dimensions["A"].width = 90
    return row


# ═══════════════════════════════════════════════════════════════════════════════
#  WORD
# ═══════════════════════════════════════════════════════════════════════════════

def create_word(answer: str) -> Path:
    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.8)
        sec.right_margin  = Cm(2.8)

    # ── Header ────────────────────────────────────────────────────────────────
    h = doc.add_heading("Moodle AI Assistant — NMIT", level=1)
    _color_run(h.runs[0], 0x1a, 0x2b, 0x8c, size=20)
    h.paragraph_format.space_after = Pt(2)

    meta = doc.add_paragraph(f"Generated: {_ts()}")
    r = meta.runs[0]
    r.font.size      = Pt(9)
    r.font.italic    = True
    r.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    meta.paragraph_format.space_after = Pt(10)

    _word_hr(doc, "1a2b8c", thick=True)
    doc.add_paragraph()

    # ── Content ───────────────────────────────────────────────────────────────
    headers, rows = _extract_table(answer)
    if headers and rows:
        non_tbl = re.sub(r"\|.*\|", "", answer, flags=re.MULTILINE).strip()
        if non_tbl:
            _word_lines(doc, non_tbl)
            doc.add_paragraph()
        _word_table(doc, headers, rows)
    else:
        _word_lines(doc, answer)

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    _word_hr(doc, "d0d5f0", thick=False)
    footer = doc.add_paragraph("NMIT Smart Campus · Moodle AI Assistant · Confidential")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.runs[0]
    r.font.size      = Pt(8)
    r.font.italic    = True
    r.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", prefix="moodle_ai_", delete=False)
    tmp.close()
    doc.save(tmp.name)
    return Path(tmp.name)


def _word_lines(doc, answer):
    for p in _parse(answer):
        t, text = p["type"], p["text"]

        if t == "blank":
            pg = doc.add_paragraph()
            pg.paragraph_format.space_after = Pt(3)
            continue

        if t == "hr":
            _word_hr(doc, "d0d5f0")
            continue

        if t == "h1":
            h = doc.add_heading(_clean_md(text), level=2)
            for run in h.runs:
                _color_run(run, 0x1a, 0x2b, 0x8c, size=14)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after  = Pt(4)
            continue

        if t in ("h2", "h3"):
            h = doc.add_heading(_clean_md(text), level=3)
            for run in h.runs:
                _color_run(run, 0x43, 0x61, 0xee, size=11)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after  = Pt(3)
            continue

        if t == "bullet":
            pg = doc.add_paragraph(style="List Bullet")
            _add_rich(pg, text)
            pg.paragraph_format.space_after = Pt(2)
            continue

        if t == "numbered":
            pg = doc.add_paragraph(style="List Number")
            _add_rich(pg, text)
            pg.paragraph_format.space_after = Pt(2)
            continue

        # Normal paragraph — preserve inline bold
        pg = doc.add_paragraph()
        pg.paragraph_format.space_after = Pt(3)
        _add_rich(pg, text)


def _add_rich(para, text):
    """Split on **bold** and add styled runs."""
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            r.bold      = True
            r.font.size = Pt(10.5)
        else:
            r = para.add_run(_clean_md(part))
            r.font.size = Pt(10.5)


def _color_run(run, r, g, b, size=None):
    run.font.color.rgb = RGBColor(r, g, b)
    if size:
        run.font.size = Pt(size)


def _word_hr(doc, hex_color, thick=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "12" if thick else "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), hex_color.lstrip("#"))
    pBdr.append(bot)
    pPr.append(pBdr)


def _word_table(doc, headers, rows):
    ncols = len(headers)
    tbl   = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h.upper()
        r = cell.paragraphs[0].runs[0]
        r.bold           = True
        r.font.size      = Pt(10)
        r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        _cell_bg(cell, "1A2B8C")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        bg = "EEF1FB" if ri % 2 else "FFFFFF"
        for ci, val in enumerate(row_data[:ncols]):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = str(val)
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            _cell_bg(cell, bg)


def _cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)